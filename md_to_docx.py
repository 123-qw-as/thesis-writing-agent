"""
Markdown to Word Converter for Thesis
将论文Markdown转换为Word格式，保留格式和结构
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_word(input_path: str, output_path: str):
    """将Markdown论文转换为Word文档"""

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # 跳过HTML注释
        if line.strip().startswith('<!--') and line.strip().endswith('-->'):
            i += 1
            continue

        # 处理分隔线
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # 处理标题
        if line.startswith('# '):
            heading = line[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            heading = line[3:].strip()
            p = doc.add_heading(heading, level=2)
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_heading(heading, level=3)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_heading(heading, level=4)
        # 处理加粗标题（关键字等）
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
            run.font.size = Pt(12)
        # 处理代码块
        elif line.startswith('```'):
            # 收集代码块内容
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            # 添加代码块
            if code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                code_text = '\n'.join(code_lines)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        # 处理列表项
        elif line.strip().startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
        elif line.strip().startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
        # 处理有序列表
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
        # 处理引用块
        elif line.startswith('>'):
            text = line[1:].strip()
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.color.rgb = RGBColor(128, 128, 128)
        # 处理普通段落
        elif line.strip():
            text = line
            p = doc.add_paragraph()

            # 解析行内格式
            parts = parse_inline_format(text)
            for part_text, bold, italic in parts:
                if part_text:
                    run = p.add_run(part_text)
                    run.bold = bold
                    run.italic = italic
                    run.font.size = Pt(12)

        i += 1

    # 添加页眉
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "基于 LangChain 框架的检索增强生成系统研究与实现"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"Word document saved to: {output_path}")


def parse_inline_format(text: str):
    """解析行内格式（加粗、斜体等）"""
    parts = []
    current_text = ''
    bold = False
    italic = False

    i = 0
    while i < len(text):
        char = text[i]

        if i + 1 < len(text) and text[i:i+2] == '**':
            if current_text:
                parts.append((current_text, bold, italic))
                current_text = ''
            bold = not bold
            i += 2
        elif text[i] == '*' and not (i > 0 and text[i-1] == '*'):
            if current_text:
                parts.append((current_text, bold, italic))
                current_text = ''
            italic = not italic
            i += 1
        elif text[i] == '`':
            if current_text:
                parts.append((current_text, bold, italic))
                current_text = ''
            end = text.find('`', i+1)
            if end != -1:
                parts.append((text[i+1:end], False, False))
                i = end + 1
            else:
                current_text += '`'
                i += 1
        else:
            current_text += char
            i += 1

    if current_text:
        parts.append((current_text, bold, italic))

    return parts if parts else [('', False, False)]


if __name__ == '__main__':
    input_file = 'output/test_thesis_optimized.md'
    output_file = 'output/thesis_rag_langchain.docx'

    markdown_to_word(input_file, output_file)
    print("Conversion completed!")