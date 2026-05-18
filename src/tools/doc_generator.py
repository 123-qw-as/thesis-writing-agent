"""
Document Generator Agent - 论文文档生成Agent
将Markdown论文转换为排版规范的Word/PDF文档
"""

import re
import os
from typing import Dict, Any, Optional, List
from dataclasses import dataclass
from datetime import datetime

from src.evaluation.rendering_auditor import RenderingAuditor


class ThesisFormatter:
    """论文格式器 - 规范Markdown格式"""

    @staticmethod
    def normalize_sections(content: str) -> str:
        """标准化章节编号"""
        lines = content.split('\n')
        result = []
        chapter_num = 0
        section_num = 0
        subsection_num = 0

        for line in lines:
            if line.startswith('## 第') and '章' in line:
                chapter_num += 1
                section_num = 0
                subsection_num = 0
                result.append(line)
            elif line.startswith('## '):
                chapter_num += 1
                section_num = 0
                subsection_num = 0
                parts = line.split('## ', 1)
                result.append(f'## 第{_to_chinese(chapter_num)}章 {parts[1]}')
            elif line.startswith('### '):
                section_num += 1
                subsection_num = 0
                parts = line.split('### ', 1)
                result.append(f'### {chapter_num}.{section_num} {parts[1]}')
            elif line.startswith('#### '):
                subsection_num += 1
                parts = line.split('#### ', 1)
                result.append(f'#### {chapter_num}.{section_num}.{subsection_num} {parts[1]}')
            else:
                result.append(line)

        return '\n'.join(result)

    @staticmethod
    def ensure_complete_structure(content: str) -> str:
        """确保论文结构完整"""
        required = [
            ('# ', '论文标题'),
            ('## 摘要', '摘要'),
            ('## 第', '章节'),
            ('## 结论', '结论'),
            ('## 参考文献', '参考文献'),
        ]

        missing = [name for pattern, name in required if not any(
            line.startswith(pattern) for line in content.split('\n')
        )]

        if missing:
            print(f'  Missing sections: {missing}')

        return content


def _to_chinese(num: int) -> str:
    """数字转中文"""
    chinese_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    if num <= 10:
        return chinese_nums[num]
    if num < 20:
        return '十' + chinese_nums[num - 10]
    return str(num)


class WordDocumentBuilder:
    """Word文档构建器 - 生成符合学术规范的Word文档"""

    def __init__(self):
        self.doc = None
        self._omml_transform = None
        self._auditor = RenderingAuditor()
        self._ensure_library()

    def _ensure_library(self):
        try:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_ORIENT
            from docx.oxml.ns import qn, nsdecls
            from docx.oxml import parse_xml
            Document  # verify import works
        except ImportError:
            import subprocess
            subprocess.run(['py', '-3', '-m', 'pip', 'install', 'python-docx'], capture_output=True)
            from docx import Document
        finally:
            from docx import Document
            from docx.shared import Pt, Inches, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            self.Document = Document
            self.Pt = Pt
            self.Inches = Inches
            self.Cm = Cm
            self.RGBColor = RGBColor
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
            self.qn = qn

    def _init_omml(self):
        """初始化LaTeX→OMML转换器"""
        if self._omml_transform is not None:
            return True
        try:
            from latex2mathml.converter import convert as latex_to_mathml
            from lxml import etree
            import os as _os
            xslt_path = _os.path.join(_os.path.dirname(__file__), '..', '..', 'MML2OMML.XSL')
            if not _os.path.exists(xslt_path):
                xslt_path = 'MML2OMML.XSL'
            if not _os.path.exists(xslt_path):
                return False
            xslt = etree.parse(xslt_path)
            self._omml_transform = (latex_to_mathml, etree.XSLT(xslt))
            return True
        except Exception:
            return False

    def _latex_to_omml(self, latex: str, display: bool = False):
        """将LaTeX转换为OMML XML元素"""
        self._auditor.record_formula_request()
        if not self._init_omml():
            self._auditor.record_fallback('formula_fallback', latex[:50], 'OMML init失败')
            self._auditor.record_formula_fallback()
            return None
        try:
            latex_to_mathml, transform = self._omml_transform
            mathml = latex_to_mathml(latex)
            from lxml import etree
            tree = etree.fromstring(mathml.encode('utf-8'))
            tree.attrib['xmlns'] = 'http://www.w3.org/1998/Math/MathML'
            if display:
                tree.attrib['display'] = 'block'
            omml = transform(tree)
            self._auditor.record_formula_omml()
            return omml.getroot()
        except Exception as e:
            self._auditor.record_fallback('formula_fallback', latex[:50], str(e)[:80])
            self._auditor.record_formula_fallback()
            return None

    def build(self, markdown_content: str, title: str = '',
             author: str = '', output_path: str = 'output/thesis.docx') -> str:
        """构建Word文档"""
        self._ensure_library()
        doc = self.Document()

        # ========== 页面设置 ==========
        section = doc.sections[0]
        section.top_margin = self.Cm(2.54)
        section.bottom_margin = self.Cm(2.54)
        section.left_margin = self.Cm(3.17)
        section.right_margin = self.Cm(3.17)

        # ========== 样式设置 ==========
        self._setup_styles(doc)

        # 解析Markdown
        lines = markdown_content.split('\n')
        self._render_content(doc, lines, title)

        # ========== 页眉页脚 ==========
        self._add_header_footer(doc, title)

        # ========== 保存 ==========
        os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
        doc.save(output_path)
        return output_path

    def _setup_styles(self, doc):
        """设置文档样式"""
        # 正文样式
        normal = doc.styles['Normal']
        normal.font.name = '宋体'
        normal.element.rPr.rFonts.set(self.qn('w:eastAsia'), '宋体')
        normal.font.size = self.Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.first_line_indent = self.Pt(24)

        # 标题样式
        for level, (size, bold) in enumerate([
            (22, True),   # Heading 1 - 论文标题
            (16, True),  # Heading 2 - 章标题
            (14, True),  # Heading 3 - 节标题
            (12, True),  # Heading 4 - 小节标题
        ], 1):
            style_name = f'Heading {level}'
            if style_name in doc.styles:
                h = doc.styles[style_name]
                h.font.name = '黑体'
                h.element.rPr.rFonts.set(self.qn('w:eastAsia'), '黑体')
                h.font.size = self.Pt(size)
                h.font.bold = bold
                h.paragraph_format.space_before = self.Pt(12)
                h.paragraph_format.space_after = self.Pt(6)
                h.paragraph_format.first_line_indent = self.Pt(0)
                h.paragraph_format.line_spacing = 1.5

    def _render_content(self, doc, lines: List[str], title: str):
        """渲染文档内容"""
        i = 0
        in_code_block = False
        code_lines = []
        title_rendered = False
        in_display_math = False
        math_lines = []

        while i < len(lines):
            line = lines[i].strip()

            if line.startswith('```'):
                if in_code_block:
                    self._render_code_block(doc, code_lines)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(lines[i].rstrip())
                i += 1
                continue

            if line.startswith('$$'):
                if in_display_math:
                    self._render_display_math(doc, math_lines)
                    math_lines = []
                    in_display_math = False
                else:
                    in_display_math = True
                i += 1
                continue

            if in_display_math:
                math_lines.append(line)
                i += 1
                continue

            # 渲染标题
            if line.startswith('# ') and not title_rendered:
                p = doc.add_heading(line[2:], level=1)
                p.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = self.Pt(36)
                p.paragraph_format.space_after = self.Pt(18)
                title_rendered = True

            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
                p.alignment = self.WD_ALIGN_PARAGRAPH.LEFT

            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
                p.alignment = self.WD_ALIGN_PARAGRAPH.LEFT

            elif line.startswith('#### '):
                p = doc.add_heading(line[5:], level=4)

            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.bold = True
                run.font.size = self.Pt(12)

            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.first_line_indent = self.Pt(0)
                run = p.add_run(text)
                run.font.size = self.Pt(12)

            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.first_line_indent = self.Pt(0)
                run = p.add_run(text)
                run.font.size = self.Pt(12)

            elif line.startswith('---') or line.startswith('***'):
                pass  # 分隔线

            elif line.startswith('!['):
                self._render_image(doc, line)

            elif line:
                self._render_paragraph(doc, lines[i].strip())

            i += 1

    def _render_paragraph(self, doc, text: str):
        """渲染段落（支持加粗、行内代码、行内OMML公式）"""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = self.Pt(24)

        # Split text by math markers and render each part
        self._render_paragraph_parts(p, text)

    def _render_paragraph_parts(self, p, text: str):
        """将段落文本分段渲染，处理行内格式和OMML公式"""
        parts = self._parse_inline_omml(text)
        for part in parts:
            ptype = part.get('type', 'text')
            if ptype == 'math':
                latex = part.get('content', '')
                omml = self._latex_to_omml(latex)
                if omml is not None:
                    p._element.append(omml)
                else:
                    readable = self._latex_to_readable(latex)
                    run = p.add_run(readable)
                    run.italic = True
                    run.font.size = self.Pt(12)
            elif ptype == 'code':
                run = p.add_run(part.get('content', ''))
                run.font.name = 'Courier New'
                run.font.size = self.Pt(12)
            elif ptype == 'bold':
                run = p.add_run(part.get('content', ''))
                run.bold = True
                run.font.size = self.Pt(12)
            elif ptype == 'italic':
                run = p.add_run(part.get('content', ''))
                run.italic = True
                run.font.size = self.Pt(12)
            else:
                run = p.add_run(part.get('content', ''))
                run.font.size = self.Pt(12)

    def _parse_inline_omml(self, text: str):
        """解析行内格式，返回部件列表（支持OMML公式）"""
        parts = []
        i = 0
        buf = ''
        bold = False
        italic = False

        while i < len(text):
            # **bold**
            if text[i:i+2] == '**':
                if buf:
                    parts.append({'type': 'bold' if bold else 'italic' if italic else 'text',
                                 'content': buf})
                bold = not bold
                buf = ''
                i += 2
                continue

            # $...$ inline math (check for $$ first)
            if text[i:i+2] == '$$':
                if buf:
                    parts.append({'type': 'bold' if bold else 'italic' if italic else 'text',
                                 'content': buf})
                j = text.find('$$', i+2)
                if j != -1:
                    parts.append({'type': 'math', 'content': text[i+2:j]})
                    i = j + 2
                else:
                    buf += '$$'
                    i += 2
                buf = ''
                continue

            if text[i] == '$':
                if buf:
                    parts.append({'type': 'bold' if bold else 'italic' if italic else 'text',
                                 'content': buf})
                j = text.find('$', i+1)
                if j != -1:
                    parts.append({'type': 'math', 'content': text[i+1:j]})
                    i = j + 1
                else:
                    buf += '$'
                    i += 1
                buf = ''
                continue

            # *italic*
            if text[i] == '*':
                if buf:
                    parts.append({'type': 'bold' if bold else 'italic' if italic else 'text',
                                 'content': buf})
                italic = not italic
                buf = ''
                i += 1
                continue

            # `code`
            if text[i] == '`':
                if buf:
                    parts.append({'type': 'bold' if bold else 'italic' if italic else 'text',
                                 'content': buf})
                j = text.find('`', i+1)
                if j != -1:
                    parts.append({'type': 'code', 'content': text[i+1:j]})
                    i = j + 1
                else:
                    buf += '`'
                    i += 1
                buf = ''
                continue

            buf += text[i]
            i += 1

        if buf:
            parts.append({'type': 'bold' if bold else 'italic' if italic else 'text',
                         'content': buf})

        return parts

    def _render_display_math(self, doc, math_lines: List[str]):
        """渲染独立显示公式（居中，使用OMML原生公式）"""
        if not math_lines:
            return
        math_text = ' '.join(math_lines)

        omml = self._latex_to_omml(math_text, display=True)
        if omml is not None:
            from lxml import etree
            MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
            p = doc.add_paragraph()
            p.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = self.Pt(12)
            p.paragraph_format.space_after = self.Pt(12)
            # Wrap in oMathPara for display math
            oMathPara = etree.SubElement(p._element, etree.QName(MATH_NS, 'oMathPara'))
            oMathPara.append(omml)
        else:
            p = doc.add_paragraph()
            p.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            readable = self._latex_to_readable(math_text)
            run = p.add_run(readable)
            run.italic = True
            run.font.size = self.Pt(11)

    def _render_image(self, doc, line: str):
        """渲染图片：![alt](path)"""
        import os, re
        match = re.match(r'!\[.*?\]\((.*?)\)', line)
        if not match:
            self._auditor.record_warning(f'图片格式无法解析: {line[:60]}')
            return
        img_path = match.group(1).strip()
        img_path = img_path.replace('\\', '/')
        self._auditor.record_image_request()

        found_path = None

        # 1. Try exact path as-is (PNG preferred for python-docx)
        if os.path.exists(img_path):
            found_path = img_path
        # 2. Try replacing extension with .png
        else:
            img_base = img_path.rsplit('.', 1)[0] if '.' in img_path else img_path
            for ext in ['.png', '.jpg', '.jpeg']:
                candidate = img_base + ext
                if os.path.exists(candidate):
                    found_path = candidate
                    break

        # 3. Fallback: search output/figures/ by basename
        if not found_path:
            base_name = os.path.basename(img_path.rsplit('.', 1)[0] if '.' in img_path else img_path)
            for ext in ['.png', '.jpg', '.jpeg']:
                candidate = 'output/figures/' + base_name + ext
                if os.path.exists(candidate):
                    found_path = candidate
                    break

        if found_path:
            # python-docx cannot embed SVG; skip SVG files
            if found_path.lower().endswith('.svg'):
                png_candidate = found_path.rsplit('.', 1)[0] + '.png'
                if os.path.exists(png_candidate):
                    found_path = png_candidate
                else:
                    self._auditor.record_fallback('image_unsupported_format', found_path, 'SVG不直接支持，需PNG')
                    base_name = os.path.basename(found_path).rsplit('.', 1)[0]
                    for ext in ['.png', '.jpg', '.jpeg']:
                        candidate = 'output/figures/' + base_name + ext
                        if os.path.exists(candidate):
                            found_path = candidate
                            break
                    else:
                        self._auditor.record_fallback('image_missing_png', found_path, '未找到对应PNG')
                        p = doc.add_paragraph()
                        run = p.add_run(f'[Image not found: {img_path}]')
                        run.italic = True
                        self._auditor.record_fallback('image_missing', img_path, '未找到任何可嵌入的图片文件')
                        return
            try:
                doc.add_picture(found_path, width=self.Inches(5.5))
                self._auditor.record_image_embedded()
                cap = doc.add_paragraph()
                cap.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.first_line_indent = self.Pt(0)
                run = cap.add_run(os.path.basename(found_path))
                run.font.size = self.Pt(9)
                run.font.color.rgb = self.RGBColor(128, 128, 128)
            except Exception as e:
                self._auditor.record_fallback('image_error', found_path, str(e)[:80])
                p = doc.add_paragraph()
                run = p.add_run(f'[Image: {os.path.basename(found_path)}]')
                run.italic = True
        else:
            self._auditor.record_fallback('image_missing', img_path, '图片文件未找到')
            p = doc.add_paragraph()
            run = p.add_run(f'[Image not found: {img_path}]')
            run.italic = True

    def _parse_inline(self, text: str):
        """解析行内格式标记（含数学公式 $...$ 和 $$...$$）"""
        parts = []
        i = 0
        buf = ''
        bold = False
        italic = False

        while i < len(text):
            # **bold** (must check before single *)
            if text[i:i+2] == '**':
                if buf:
                    parts.append((buf, bold, italic, False))
                bold = not bold
                buf = ''
                i += 2
                continue

            # $$ display math $$
            if text[i:i+2] == '$$':
                if buf:
                    parts.append((buf, bold, italic, False))
                j = text.find('$$', i+2)
                if j != -1:
                    math_content = self._latex_to_readable(text[i+2:j])
                    parts.append((math_content, False, True, False))
                    i = j + 2
                else:
                    buf += '$$'
                    i += 2
                continue

            # $ inline math $
            if text[i] == '$':
                if buf:
                    parts.append((buf, bold, italic, False))
                j = text.find('$', i+1)
                if j != -1:
                    math_content = self._latex_to_readable(text[i+1:j])
                    parts.append((math_content, False, True, False))
                    i = j + 1
                else:
                    buf += '$'
                    i += 1
                continue

            # *italic* (single star)
            if text[i] == '*':
                if buf:
                    parts.append((buf, bold, italic, False))
                italic = not italic
                buf = ''
                i += 1
                continue

            # `code`
            if text[i] == '`':
                if buf:
                    parts.append((buf, bold, italic, False))
                j = text.find('`', i+1)
                if j != -1:
                    parts.append((text[i+1:j], False, False, True))
                    i = j + 1
                else:
                    buf += '`'
                    i += 1
                continue

            # normal character
            buf += text[i]
            i += 1

        if buf:
            parts.append((buf, bold, italic, False))

        return parts

    def _latex_to_readable(self, latex: str) -> str:
        """将LaTeX数学表达式转换为可读文本"""
        import re
        text = latex.strip()

        # Handle \frac{a}{b}
        while '\\frac' in text:
            text = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', text, count=1)

        # Handle \sqrt[n]{x}
        text = re.sub(r'\\sqrt\[([^}]*)\]\{([^}]*)\}', r'sqrt[\1](\2)', text)
        text = re.sub(r'\\sqrt\{([^}]*)\}', r'sqrt(\1)', text)

        # Handle \left, \right
        text = re.sub(r'\\left', '', text)
        text = re.sub(r'\\right', '', text)

        # Handle superscripts and subscripts
        text = re.sub(r'\^\{([^}]*)\}', r'^{\1}', text)
        text = re.sub(r'_\{([^}]*)\}', r'_{\1}', text)

        # Handle \text, \mathrm, \mathbf, \mathcal, \operatorname
        text = re.sub(r'\\text\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', text)
        text = re.sub(r'\\operatorname\{([^}]*)\}', r'\1', text)

        # Handle \mathbb
        text = re.sub(r'\\mathbb\{([^}]*)\}', r'\1', text)

        # Common math symbols
        symbols = [
            (r'\\in', ' in '),
            (r'\\otimes', ' (x) '),
            (r'\\times', ' x '),
            (r'\\cdot', '.'),
            (r'\\pm', '+/-'),
            (r'\\infty', 'inf'),
            (r'\\approx', '~='),
            (r'\\neq', '!='),
            (r'\\leq', '<='),
            (r'\\geq', '>='),
            (r'\\rightarrow', '->'),
            (r'\\leftarrow', '<-'),
            (r'\\Rightarrow', '=>'),
            (r'\\Leftarrow', '<='),
            (r'\\partial', 'd'),
            (r'\\nabla', 'grad'),
            (r'\\sum', 'SUM'),
            (r'\\prod', 'PROD'),
            (r'\\int', 'INT'),
        ]
        for pattern, replacement in symbols:
            text = re.sub(pattern, replacement, text)

        # Greek letters
        greek = ['alpha','beta','gamma','delta','epsilon','zeta','eta','theta',
                 'iota','kappa','lambda','mu','nu','xi','omicron','pi','rho',
                 'sigma','tau','upsilon','phi','chi','psi','omega',
                 'Gamma','Delta','Theta','Lambda','Xi','Pi','Sigma','Phi','Psi','Omega']
        for letter in sorted(greek, key=len, reverse=True):
            text = re.sub(f'\\\\{letter}', letter, text)

        # Named functions
        for func in ['sin', 'cos', 'tan', 'log', 'ln', 'exp', 'max', 'min', 'det', 'rank']:
            text = re.sub(f'\\\\{func}', func, text)

        # Clean up remaining backslash commands
        text = re.sub(r'\\[a-zA-Z]+', '', text)

        # Clean up braces
        text = text.replace('{', '').replace('}', '')

        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        return text

    def _render_code_block(self, doc, code_lines: List[str]):
        """渲染代码块"""
        if not code_lines:
            return
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = self.Cm(1)
        p.paragraph_format.first_line_indent = self.Pt(0)
        p.paragraph_format.space_before = self.Pt(6)
        p.paragraph_format.space_after = self.Pt(6)
        run = p.add_run('\n'.join(code_lines))
        run.font.name = 'Courier New'
        run.font.size = self.Pt(9)

    def _add_header_footer(self, doc, title: str):
        """添加页眉页脚"""
        for section in doc.sections:
            # 页眉
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0]
            hp.text = title[:50] if title else ''
            hp.alignment = self.WD_ALIGN_PARAGRAPH.CENTER
            hp.style.font.size = self.Pt(10)

            # 页脚 - 页码
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]
            fp.alignment = self.WD_ALIGN_PARAGRAPH.CENTER

            run = fp.add_run()
            fldChar1 = run._element.makeelement(self.qn('w:fldChar'), {self.qn('w:fldCharType'): 'begin'})
            run._element.append(fldChar1)

            run2 = fp.add_run()
            instrText = run2._element.makeelement(self.qn('w:instrText'), {})
            instrText.text = ' PAGE '
            run2._element.append(instrText)

            run3 = fp.add_run()
            fldChar2 = run3._element.makeelement(self.qn('w:fldChar'), {self.qn('w:fldCharType'): 'end'})
            run3._element.append(fldChar2)


class ThesisDocumentGenerator:
    """论文文档生成器 - 整合所有功能"""

    def __init__(self):
        self.formatter = ThesisFormatter()
        self.word_builder = WordDocumentBuilder()

    def generate_docx(
        self,
        content: str,
        title: str = '',
        author: str = '',
        output_path: str = 'output/thesis.docx',
        normalize: bool = True
    ) -> str:
        """
        生成Word文档

        Args:
            content: Markdown格式论文内容
            title: 论文标题
            author: 作者
            output_path: 输出路径
            normalize: 是否自动规范化格式

        Returns:
            输出文件路径
        """
        if normalize:
            content = self.formatter.normalize_sections(content)
            content = self.formatter.ensure_complete_structure(content)

        return self.word_builder.build(content, title, output_path=output_path)

    def generate_pdf(
        self,
        content: str,
        title: str = '',
        author: str = '',
        output_path: str = 'output/thesis.pdf'
    ) -> str:
        """
        生成PDF文档
        先生成Word再转换PDF
        """
        docx_path = output_path.replace('.pdf', '.docx')
        self.generate_docx(content, title, author, docx_path)

        try:
            import subprocess
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'pdf', docx_path, '--outdir',
                 os.path.dirname(output_path)],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0:
                return output_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

        return docx_path


# 快捷函数
def generate_thesis_docx(
    markdown_content: str,
    title: str = '',
    author: str = '',
    output_path: str = 'output/thesis.docx'
) -> str:
    """快捷生成Word文档"""
    generator = ThesisDocumentGenerator()
    return generator.generate_docx(markdown_content, title, author, output_path)