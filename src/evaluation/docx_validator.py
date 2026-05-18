"""
DOCX Validator - 编译后Word文档质量验证器
打开生成的DOCX文件，检查内容完整性、公式渲染、图片嵌入等
"""
import os
import re
import zipfile
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field

from src.evaluation.rendering_auditor import RenderingAuditor, RenderingReport
from src.evaluation.rubrics.docx_rubric import DocxRubric
from src.evaluation.judge import LLMJudge


@dataclass
class DocxAnalysis:
    """DOCX分析结果"""
    file_size: int = 0
    paragraph_count: int = 0
    heading_count: int = 0
    image_count: int = 0
    inline_shape_count: int = 0
    omml_formula_count: int = 0
    omml_display_count: int = 0
    has_header: bool = False
    has_footer: bool = False
    has_page_number: bool = False
    margins: dict = field(default_factory=dict)
    line_spacing: float = 0
    degradation_markers: List[str] = field(default_factory=list)
    styles_used: set = field(default_factory=set)
    errors: List[str] = field(default_factory=list)


class DocxValidator:
    """DOCX文档后验器"""

    def __init__(self, llm=None):
        self.llm = llm
        self.auditor = RenderingAuditor()

    def validate(self, docx_path: str, rendering_report: RenderingReport = None) -> DocxAnalysis:
        """
        验证生成的DOCX文件质量
        
        Args:
            docx_path: DOCX文件路径
            rendering_report: 渲染审计报告（可选）
            
        Returns:
            DocxAnalysis: DOCX分析结果
        """
        analysis = DocxAnalysis()
        analysis.file_size = os.path.getsize(docx_path) if os.path.exists(docx_path) else 0

        if not os.path.exists(docx_path):
            analysis.errors.append(f'文件不存在: {docx_path}')
            return analysis

        try:
            from docx import Document
            doc = Document(docx_path)
        except Exception as e:
            analysis.errors.append(f'无法打开DOCX: {e}')
            return analysis

        # 1. 段落与标题
        analysis.paragraph_count = len(doc.paragraphs)
        analysis.heading_count = len([
            p for p in doc.paragraphs
            if p.style and 'Heading' in (p.style.name or '')
        ])

        # 2. 图片
        analysis.inline_shape_count = len(doc.inline_shapes)

        # 3. 样式
        for p in doc.paragraphs:
            if p.style:
                analysis.styles_used.add(p.style.name)

        # 4. 页眉页脚
        for section in doc.sections:
            if section.header.paragraphs and any(p.text.strip() for p in section.header.paragraphs):
                analysis.has_header = True
            if section.footer.paragraphs:
                analysis.has_footer = True
                for p in section.footer.paragraphs:
                    if 'PAGE' in p.text.upper() or any(
                        r.text and ('PAGE' in r.text.upper() or '1' in r.text)
                        for r in p.runs
                    ):
                        analysis.has_page_number = True

            # 5. 边距
            analysis.margins = {
                'top': section.top_margin,
                'bottom': section.bottom_margin,
                'left': section.left_margin,
                'right': section.right_margin,
            }

            # 6. 行距（取第一个正文段落）
            for p in doc.paragraphs:
                if p.style and p.style.name == 'Normal':
                    analysis.line_spacing = p.paragraph_format.line_spacing or 0
                    break

        # 7. 解析XML，检查OMML和退化标记
        try:
            with zipfile.ZipFile(docx_path) as zf:
                doc_xml = zf.read('word/document.xml').decode('utf-8', errors='replace')
                analysis.omml_formula_count = doc_xml.count('m:oMath')
                analysis.omml_display_count = doc_xml.count('m:oMathPara')
                analysis.image_count = len([n for n in zf.namelist() if n.startswith('word/media/')])

                # 8. 退化标记搜索
                for marker in ['[Image not found]', '[Image:', '[Image error]', '\\mathbb', '\\times', '\\frac']:
                    count = doc_xml.count(marker)
                    if count > 0:
                        analysis.degradation_markers.append(f'{marker}({count}处)')
        except Exception as e:
            analysis.errors.append(f'XML解析失败: {e}')

        # 9. 合并渲染报告
        if rendering_report:
            analysis.inline_shape_count += rendering_report.images_embedded
            if rendering_report.formulas_fallback > 0:
                analysis.degradation_markers.append(
                    f'公式退化({rendering_report.formulas_fallback}处)'
                )

        return analysis

    def evaluate(self, docx_path: str) -> Dict[str, Any]:
        """
        对DOCX执行完整评估
        
        Returns:
            {
                'analysis': DocxAnalysis,
                'rubric_score': float,
                'passed': bool,
                'issues': [str],
                'suggestions': [str]
            }
        """
        rendering_report = self.auditor.get_report()
        analysis = self.validate(docx_path, rendering_report)

        # 执行Rubric评估
        from src.evaluation.judge import LLMJudge
        judge = LLMJudge(self.llm) if self.llm else None
        rubric = DocxRubric(judge)
        
        analysis_dict = {
            'has_header': analysis.has_header,
            'has_footer': analysis.has_footer,
            'heading_count': analysis.heading_count,
            'paragraph_count': analysis.paragraph_count,
            'margins': analysis.margins,
            'line_spacing': analysis.line_spacing,
        }
        eval_result = rubric.evaluate_from_report(rendering_report, docx_path, analysis_dict)

        # 生成问题和建议
        issues = []
        suggestions = []

        if rendering_report.formulas_fallback > 0:
            issues.append(f'{rendering_report.formulas_fallback}个公式退化文本')
            suggestions.append('检查 latex2mathml 和 MML2OMML.XSL 是否正常安装')

        if analysis.inline_shape_count == 0 and analysis.image_count == 0:
            issues.append('文档中无任何图片')
            suggestions.append('检查图片文件路径和 add_picture 调用')

        if analysis.degradation_markers:
            issues.append(f'退化标记: {", ".join(analysis.degradation_markers[:3])}')
            suggestions.append('修复图片路径或公式渲染配置')

        if analysis.file_size < 20000:
            issues.append(f'文件过小({analysis.file_size}字节)')
            suggestions.append('检查是否有图片/内容被遗漏')

        return {
            'analysis': analysis,
            'rubric_score': eval_result.overall_score,
            'passed': eval_result.passes(),
            'dimensions': [
                {'name': d.name, 'score': d.score, 'reasoning': d.reasoning}
                for d in eval_result.dimensions
            ],
            'issues': issues,
            'suggestions': suggestions,
            'rendering_report': rendering_report.summary(),
        }


def validate_docx(docx_path: str, llm=None) -> Dict[str, Any]:
    """快捷验证函数"""
    validator = DocxValidator(llm)
    return validator.evaluate(docx_path)